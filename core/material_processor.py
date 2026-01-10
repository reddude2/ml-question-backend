import os
from typing import Dict, List, Optional
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import pytesseract  # Optional: for image OCR

class MaterialProcessor:
    """
    Process uploaded learning materials
    Extract text and analyze content
    """
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def process_file(self, file_path: str) -> Dict:
        """
        Process uploaded file and extract information
        
        Args:
            file_path: Path to uploaded file
            
        Returns:
            Dict with extracted_text, file_info, topics
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file info
        file_size = os.path.getsize(file_path)
        file_name = os.path.basename(file_path)
        file_ext = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Extract text based on file type
        extracted_text = self.extract_text(file_path, file_ext)
        
        # Validate text
        if not extracted_text or len(extracted_text.strip()) < 100:
            raise ValueError("Extracted text too short (minimum 100 characters)")
        
        # Analyze content (simple version)
        word_count = len(extracted_text.split())
        char_count = len(extracted_text)
        
        return {
            'extracted_text': extracted_text,
            'file_info': {
                'name': file_name,
                'size': file_size,
                'type': file_ext,
                'word_count': word_count,
                'char_count': char_count
            },
            'topics': []  # Will be filled by AI later
        }
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from file based on type
        
        Args:
            file_path: Path to file
            file_type: File extension (pdf, docx, txt)
            
        Returns:
            Extracted text
        """
        if file_type == 'pdf':
            return self._extract_from_pdf(file_path)
        elif file_type == 'docx':
            return self._extract_from_docx(file_path)
        elif file_type == 'txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting PDF: {e}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {e}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read().strip()
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any encoding")
        except Exception as e:
            raise Exception(f"Error extracting TXT: {e}")
    
    def analyze_topics(self, text: str, max_length: int = 3000) -> List[str]:
        """
        Analyze text and identify main topics
        (Simple keyword-based version, can be enhanced with AI)
        
        Args:
            text: Text to analyze
            max_length: Maximum text length to analyze
            
        Returns:
            List of identified topics
        """
        # Use first N characters for analysis
        sample = text[:max_length]
        
        # Simple keyword extraction (can be replaced with AI)
        # For now, return empty list - will be filled by Gemini
        return []
    
    def validate_material_length(self, text: str, min_length: int = 100, max_length: int = 50000) -> bool:
        """
        Validate if material text is within acceptable length
        
        Args:
            text: Material text
            min_length: Minimum required length
            max_length: Maximum allowed length
            
        Returns:
            True if valid, False otherwise
        """
        text_length = len(text)
        return min_length <= text_length <= max_length
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]
        
        # Join with single newline
        cleaned = '\n'.join(lines)
        
        # Remove multiple spaces
        import re
        cleaned = re.sub(r' +', ' ', cleaned)
        
        return cleaned


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def process_uploaded_material(file_path: str) -> Dict:
    """
    Convenience function to process a material file
    
    Args:
        file_path: Path to uploaded file
        
    Returns:
        Processed material data
    """
    processor = MaterialProcessor()
    return processor.process_file(file_path)


# ============================================================================
# TESTING
# ============================================================================

if __name__ == "__main__":
    # Test material processor
    import sys
    
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        
        print(f"\n🧪 Testing Material Processor")
        print(f"File: {test_file}\n")
        
        try:
            processor = MaterialProcessor()
            result = processor.process_file(test_file)
            
            print("✅ Processing successful!")
            print(f"\nFile Info:")
            print(f"  Name: {result['file_info']['name']}")
            print(f"  Type: {result['file_info']['type']}")
            print(f"  Size: {result['file_info']['size']} bytes")
            print(f"  Words: {result['file_info']['word_count']}")
            print(f"\nExtracted Text Preview:")
            print(result['extracted_text'][:300] + "...")
            
        except Exception as e:
            print(f"❌ Error: {e}")
    else:
        print("Usage: python material_processor.py <file_path>")